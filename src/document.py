# 本文件接受列表并以此生成 Word 格式的文档。

from docx import Document

# generateDoc() 函数将翻译结果字符串函数写入 python-docx 库提供的 Document 成员
# 并返回。这里源语言的 source 和目标语言的 results 列表是分开的，source 不包含于
# results 列表中。
def generateDoc(source, results, source_bold=True):
    doc = Document()

    # 根据 source_bold 参数写入源字符串并决定它是否加粗。
    source_paragraph = doc.add_paragraph()
    source_paragraph.add_run(source).bold = source_bold

    # 逐 results 各元素写入翻译结果。
    for result in results:
        doc.add_paragraph(result)

    return doc